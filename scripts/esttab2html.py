#!/usr/bin/env python3
"""
esttab2html.py — Convert esttab CSV (plain) output to HTML + docx.

不依赖 pandoc：HTML 与 docx 均由 Python 直接生成。

Usage:
    python esttab2html.py output/tables/main_table.csv
    python esttab2html.py output/tables/main_table.csv --title "Table 2: 基准回归"
    python esttab2html.py output/tables/main_table.csv --output-dir output/tables/
    python esttab2html.py output/tables/main_table.csv --note "省份层面聚类稳健标准误"

Requires:
    python-docx (pip install python-docx)

Input:
    esttab CSV，必须带 plain 选项（否则单元格被 ="" 包裹，解析报错）。

Output:
    output/tables/main_table.html — HTML（Obsidian 预览用，inline 三线表样式）
    output/tables/main_table.docx — Word（CSSCI 投稿用，宋体 10pt 三线表）
"""

import argparse
import csv
import html
import os
import re
import sys

SE_PATTERN = re.compile(r"^\(.*\)$")
SUP_PATTERN = re.compile(r"\$\^(.*?)\$")

DEFAULT_NOTE = "括号内为标准误；* p<0.10, ** p<0.05, *** p<0.01"


def die(msg):
    print(f"错误: {msg}", file=sys.stderr)
    sys.exit(1)


def parse_esttab_csv(csv_path):
    """
    解析 esttab plain CSV。

    返回 dict:
        header_rows: list[list[str]]  表头行（depvar 名 + mtitles）
        coef_rows:   list[dict]       {label, cells, se: [str] | None}
        stat_rows:   list[list[str]]  底部统计行（FE 标注、N、Adj. R^2 等）
    """
    with open(csv_path, newline="", encoding="utf-8-sig") as f:
        rows = [row for row in csv.reader(f) if any(cell.strip() for cell in row)]

    if len(rows) < 2:
        die(f"CSV 内容过少（{len(rows)} 行），不是有效的 esttab 输出: {csv_path}")

    n_cols = max(len(r) for r in rows)
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    if n_cols < 2:
        die("CSV 只有一列，没有回归结果列（检查 esttab 是否加了 plain 选项）")

    if any('=""' in cell for row in rows for cell in row):
        die("检测到 =\"\" 包裹：esttab 输出缺少 plain 选项，请重新导出")

    # 表头：首列连续为空的行（depvar 名行 + mtitles 行，通常 1-2 行）。
    # 正文从第一个首列非空的行开始（系数标签行），不能用括号特征判断，
    # 因为 mtitles 行 ",(1),(2),(3)" 也全是括号。
    header_rows = []
    i = 0
    while i < len(rows) and rows[i][0].strip() == "":
        header_rows.append(rows[i])
        i += 1
    if not header_rows:
        die("缺少表头行（第一行应为 depvar 名 / 模型编号，首列为空）")

    body = rows[i:]
    coef_rows = []
    stat_rows = []

    def followed_by_se(idx):
        return idx + 1 < len(body) and is_se_row(body[idx + 1])

    j = 0
    while j < len(body):
        row = body[j]
        if is_se_row(row):
            # 游离的 SE 行：附到最近一个系数行
            if coef_rows and coef_rows[-1]["se"] is None:
                coef_rows[-1]["se"] = row[1:]
            else:
                die(f"第 {j + len(header_rows) + 1} 行: SE 行没有对应的系数行")
            j += 1
            continue
        label, cells = row[0].strip(), row[1:]
        if followed_by_se(j):
            coef_rows.append({"label": label, "cells": cells, "se": body[j + 1][1:]})
            j += 2
        else:
            stat_rows.append([label] + cells)
            j += 1

    if not coef_rows and not stat_rows:
        die("未解析到任何数据行，请检查 CSV 内容")

    return {"header_rows": header_rows, "coef_rows": coef_rows, "stat_rows": stat_rows}


def is_se_row(row):
    """SE 行：首列为空，且所有非空单元格都是 (…) 形式。"""
    if row[0].strip() != "":
        return False
    values = [c.strip() for c in row[1:] if c.strip()]
    return bool(values) and all(SE_PATTERN.match(v) for v in values)


def fmt_cell(text, keep_empty=False):
    """HTML 转义 + R$^2$ → R<sup>2</sup>；空单元格默认渲染为 —（结构性空位除外）。"""
    s = text.strip()
    if not s:
        return "" if keep_empty else "—"
    return SUP_PATTERN.sub(r"<sup>\1</sup>", html.escape(s))


# ---------------------------------------------------------------- HTML

TABLE_STYLE = (
    "border-collapse:collapse; border-top:1.5pt solid #000; "
    "border-bottom:1.5pt solid #000; font-family:SimSun,serif; font-size:10pt;"
)
TH_STYLE = "padding:2px 8px; font-weight:normal;"
TD_STYLE = "padding:2px 8px;"
SE_STYLE = "padding:2px 8px; color:#444;"


def render_html(parsed, title, note=None):
    n_cols = len(parsed["header_rows"][0])

    def cell_tag(tag, text, first, extra_style=""):
        align = "text-align:left;" if first else "text-align:center;"
        # 首列为标签位（表头左上角、SE 行前缀），结构性空位不填 —
        return f'<{tag} style="{align}{extra_style}">{fmt_cell(text, keep_empty=first)}</{tag}>'

    parts = []
    if title:
        parts.append(
            f'<p style="text-align:center; font-family:SimSun,serif; font-size:10.5pt;">'
            f"<strong>{html.escape(title)}</strong></p>"
        )
    parts.append(f'<table style="{TABLE_STYLE}">')

    # thead：最后一行表头下方加中线（midrule）
    parts.append("<thead>")
    for k, hr in enumerate(parsed["header_rows"]):
        last = k == len(parsed["header_rows"]) - 1
        style = TH_STYLE + ("border-bottom:1pt solid #000;" if last else "")
        cells = "".join(
            cell_tag("th", c, first=(idx == 0), extra_style=style)
            for idx, c in enumerate(hr)
        )
        parts.append(f"<tr>{cells}</tr>")
    parts.append("</thead>")

    parts.append("<tbody>")
    for cr in parsed["coef_rows"]:
        cells = "".join(
            cell_tag("td", c, first=(idx == 0), extra_style=TD_STYLE)
            for idx, c in enumerate([cr["label"]] + cr["cells"])
        )
        parts.append(f"<tr>{cells}</tr>")
        if cr["se"]:
            se_cells = "".join(
                cell_tag("td", c, first=(idx == 0), extra_style=SE_STYLE)
                for idx, c in enumerate([""] + cr["se"])
            )
            parts.append(f"<tr>{se_cells}</tr>")
    for sr in parsed["stat_rows"]:
        cells = "".join(
            cell_tag("td", c, first=(idx == 0), extra_style=TD_STYLE)
            for idx, c in enumerate(sr)
        )
        parts.append(f"<tr>{cells}</tr>")
    parts.append("</tbody>")
    parts.append("</table>")
    if note and note.strip():
        parts.append(
            '<p style="font-family:SimSun,serif; font-size:9pt;">'
            f"注：{html.escape(note)}。</p>"
        )
    return "<!-- Generated by esttab2html.py -->\n" + "\n".join(parts) + "\n"


# ---------------------------------------------------------------- docx

def _set_cell_border(cell, edge, size_eighths_pt):
    """给单元格某条边加边框（size 单位：1/8 pt）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size_eighths_pt))
    el.set(qn("w:color"), "000000")
    tcBorders.append(el)


def _set_table_borders(table):
    """三线表：仅顶线、底线，去全部网格线（中线在表头行单元格上加）。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")  # 1.5pt
        el.set(qn("w:color"), "000000")
        tblBorders.append(el)
    for edge in ("left", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _write_cell(cell, text, bold=False, first=False):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT if first else WD_ALIGN_PARAGRAPH.CENTER
    text = text.strip()
    if not text and not first:
        # 空白单元格填 —（首列为标签位，结构性空位不填）
        text = "—"
    # R$^2$ 等标记 → Word 上标 run
    for k, seg in enumerate(SUP_PATTERN.split(text)):
        if seg == "":
            continue
        run = para.add_run(seg)
        run.bold = bold
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if k % 2 == 1:
            run.font.superscript = True


def render_docx(parsed, title, docx_path, note=None):
    try:
        from docx import Document
    except ImportError:
        die("python-docx 未安装：pip install python-docx")

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    n_cols = len(parsed["header_rows"][0])
    n_body_rows = len(parsed["coef_rows"]) + sum(
        1 for c in parsed["coef_rows"] if c["se"]
    ) + len(parsed["stat_rows"])
    n_rows = len(parsed["header_rows"]) + n_body_rows

    table = doc.add_table(rows=n_rows, cols=n_cols)
    _set_table_borders(table)

    r = 0
    for k, hr in enumerate(parsed["header_rows"]):
        for idx, text in enumerate(hr):
            _write_cell(table.cell(r, idx), text, bold=False, first=(idx == 0))
        if k == len(parsed["header_rows"]) - 1:
            # 中线（midrule）：表头最后一行单元格的下边框
            for idx in range(n_cols):
                _set_cell_border(table.cell(r, idx), "bottom", 8)  # 1pt
        r += 1

    for cr in parsed["coef_rows"]:
        for idx, text in enumerate([cr["label"]] + cr["cells"]):
            _write_cell(table.cell(r, idx), text, first=(idx == 0))
        r += 1
        if cr["se"]:
            for idx, text in enumerate([""] + cr["se"]):
                _write_cell(table.cell(r, idx), text, first=(idx == 0))
            r += 1

    for sr in parsed["stat_rows"]:
        for idx, text in enumerate(sr):
            _write_cell(table.cell(r, idx), text, first=(idx == 0))
        r += 1

    if note and note.strip():
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(f"注：{note}。")
        note_run.font.size = Pt(9)
        note_run.font.name = "宋体"
        note_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    doc.save(docx_path)


# ---------------------------------------------------------------- main

def main():
    parser = argparse.ArgumentParser(
        description="Convert esttab CSV (plain) table to HTML + docx（不依赖 pandoc）"
    )
    parser.add_argument("input", help="Path to esttab .csv file (plain)")
    parser.add_argument("--title", "-t", help="Table title")
    parser.add_argument(
        "--note", "-n",
        help=f"表尾注脚（默认: {DEFAULT_NOTE}；纯数据表如 tabstat 不传则不输出注脚）",
    )
    parser.add_argument("--output-dir", "-d", help="Output directory (default: same as input)")

    args = parser.parse_args()

    csv_path = args.input
    if not os.path.exists(csv_path):
        die(f"找不到文件 {csv_path}")
    if not csv_path.lower().endswith(".csv"):
        die(f"输入必须是 esttab CSV 文件（plain 格式）: {csv_path}")

    parsed = parse_esttab_csv(csv_path)

    base = os.path.splitext(csv_path)[0]
    html_path = f"{base}.html"
    docx_path = f"{base}.docx"
    if args.output_dir:
        os.makedirs(args.output_dir, exist_ok=True)
        basename = os.path.splitext(os.path.basename(csv_path))[0]
        html_path = os.path.join(args.output_dir, f"{basename}.html")
        docx_path = os.path.join(args.output_dir, f"{basename}.docx")

    # 注脚：显式传则用传值；未传且为回归表则用默认；未传且纯数据表（tabstat）则不输出
    note = args.note
    if note is None and parsed["coef_rows"]:
        note = DEFAULT_NOTE

    with open(html_path, "w", encoding="utf-8") as f:
        f.write(render_html(parsed, args.title, note=note))
    print(f"✅ {html_path}")

    render_docx(parsed, args.title, docx_path, note=note)
    print(f"✅ {docx_path}")

    print(f"\n✅ 转换完成: {html_path} + {docx_path}")


if __name__ == "__main__":
    main()
