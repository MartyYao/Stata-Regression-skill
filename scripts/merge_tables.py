#!/usr/bin/env python3
"""
merge_tables.py — 合并 output/tables/ 下所有 .docx 为一个附录文档。

每张表前加粗体标题（读同名 .html 的 <strong> 标签，否则从文件名推导），
表间加分页符。

Usage:
    python scripts/merge_tables.py output/tables/ --output output/附录-实证表格.docx

Requires:
    python-docx (pip install python-docx)
"""

import argparse
import os
import re
import sys
from copy import deepcopy

TABLE_NUM_PATTERN = re.compile(r"table(\d+)", re.IGNORECASE)
HTML_STRONG_PATTERN = re.compile(r"<strong>(.*?)</strong>", re.IGNORECASE | re.DOTALL)
TAG_PATTERN = re.compile(r"<[^>]+>")


def die(msg):
    print(f"错误: {msg}", file=sys.stderr)
    sys.exit(1)


def _insert_before_sectpr(doc, element):
    """在 sectPr 之前插入元素（避免 OOXML 非法结构）。"""
    from docx.oxml.ns import qn
    body = doc._body._element
    sectPr = body.find(qn("w:sectPr"))
    if sectPr is not None:
        sectPr.addprevious(element)
    else:
        body.append(element)


def table_sort_key(filename):
    m = TABLE_NUM_PATTERN.search(filename)
    return (0, int(m.group(1)), filename) if m else (1, 0, filename)


def title_from_html(html_path):
    if not os.path.exists(html_path):
        return None
    with open(html_path, encoding="utf-8") as f:
        content = f.read()
    m = HTML_STRONG_PATTERN.search(content)
    if m:
        title = TAG_PATTERN.sub("", m.group(1)).strip()
        if title:
            return title
    print(f"⚠️ 未从 HTML 提取到标题（fallback 到文件名）: {html_path}")
    return None


def title_from_filename(basename):
    m = TABLE_NUM_PATTERN.search(basename)
    if not m:
        return basename.replace("_", " ")
    desc = TABLE_NUM_PATTERN.sub("", basename).strip("_")
    desc = re.sub(r"\bv(\d+)\b", lambda x: f"V{x.group(1)}", desc, flags=re.IGNORECASE)
    desc = desc.replace("_", " ")
    title = f"Table {int(m.group(1))}"
    return f"{title}: {desc}" if desc else title


def add_title(doc, text):
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def merge(tables_dir, output_path):
    try:
        from docx import Document
    except ImportError:
        die("python-docx 未安装：pip install python-docx")

    docx_files = sorted(
        (f for f in os.listdir(tables_dir) if f.lower().endswith(".docx")),
        key=table_sort_key,
    )
    out_name = os.path.basename(output_path)
    docx_files = [f for f in docx_files if f != out_name]
    if not docx_files:
        die(f"{tables_dir} 下没有 .docx 文件")

    doc = Document()
    added = 0
    for fname in docx_files:
        path = os.path.join(tables_dir, fname)
        base = os.path.splitext(fname)[0]
        title = title_from_html(os.path.join(tables_dir, f"{base}.html")) or title_from_filename(base)

        src = Document(path)
        if not src.tables:
            print(f"⚠️ 跳过（无表格）: {fname}")
            continue
        if added > 0:
            doc.add_page_break()
        add_title(doc, title)
        # 复制表格（使用 python-docx 的内部插入机制，避免 append 到 sectPr 之后）
        for tbl in src.tables:
            _insert_before_sectpr(doc, deepcopy(tbl._tbl))
        # 复制表注段落（保留强制规范要求的注脚）
        for para in src.paragraphs:
            text = para.text.strip()
            if text.startswith("注："):
                _insert_before_sectpr(doc, deepcopy(para._element))
        print(f"✅ {fname} → {title}")
        added += 1

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(
        description="合并 output/tables/ 下所有 .docx 为一个附录文档（按表号排序，表间分页）"
    )
    parser.add_argument("tables_dir", help="表格目录（如 output/tables/）")
    parser.add_argument(
        "--output", "-o", default=None,
        help="输出 docx 路径（默认: <tables_dir>/../附录-实证表格.docx）"
    )
    args = parser.parse_args()

    tables_dir = args.tables_dir.rstrip("/")
    output_path = args.output or os.path.join(os.path.dirname(tables_dir), "附录-实证表格.docx")
    merge(tables_dir, output_path)
    print(f"\n✅ 合并完成: {output_path}")


if __name__ == "__main__":
    main()
